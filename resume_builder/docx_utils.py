from __future__ import annotations

from collections.abc import Sequence

from docx.document import Document as DocumentType
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.styles.style import ParagraphStyle

from resume_builder.models import ResumeMeta
from resume_builder.theme import DEFAULT_THEME, ResumeTheme


def hex_color(value: str) -> RGBColor:
    value = value.strip().lstrip("#")
    if len(value) != 6:
        raise ValueError(f"Expected 6-digit hex color, got {value!r}")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def set_font(
    run,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    name: str | None = None,
    size: float | None = None,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor | None = None,
) -> None:
    name = name or theme.font_family
    run.font.name = name
    run.font.bold = bold
    run.font.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_style_font(
    style: ParagraphStyle,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    name: str | None = None,
    size: float = 11,
    bold: bool = False,
    color: RGBColor | None = None,
    italic: bool = False,
) -> None:
    name = name or theme.font_family
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color


def configure_paragraph_style(
    style: ParagraphStyle,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    size: float,
    bold: bool = False,
    color: RGBColor | None = None,
    italic: bool = False,
    before: float = 0,
    after: float = 0,
    line_spacing: float = 1.2,
    left_indent: Inches | None = None,
    first_line_indent: Inches | None = None,
    alignment: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    set_style_font(style, theme, size=size, bold=bold, color=color, italic=italic)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = line_spacing
    if left_indent is not None:
        style.paragraph_format.left_indent = left_indent
    if first_line_indent is not None:
        style.paragraph_format.first_line_indent = first_line_indent
    if alignment is not None:
        style.paragraph_format.alignment = alignment


def set_paragraph_spacing(
    paragraph, *, before: float = 0, after: float = 0, line_spacing: float = 1.2
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing


def set_keep_with_next(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    ppr.append(keep)


def add_bottom_border(
    paragraph,
    color: str = DEFAULT_THEME.light_grey,
    size: str = DEFAULT_THEME.section_border_size,
) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    p_bdr = ppr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        ppr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def clear_document(document: DocumentType) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def add_entry_heading(
    document: DocumentType,
    heading: str,
    date: str,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    style_name: str = "Heading 2",
    left_indent: float = 0,
    align_date: bool = True,
) -> None:
    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.left_indent = Inches(left_indent)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(theme.page_width - (2 * theme.margin)),
        WD_TAB_ALIGNMENT.RIGHT,
    )

    heading_run = paragraph.add_run(heading)
    set_font(
        heading_run,
        theme,
        size=theme.role_heading_size,
        bold=True,
        color=hex_color(theme.primary_text_color),
    )
    if align_date:
        paragraph.add_run("\t")
    else:
        paragraph.add_run(" | ")
    date_run = paragraph.add_run(date.replace(" ", "\N{NO-BREAK SPACE}"))
    set_font(
        date_run,
        theme,
        size=theme.date_size,
        color=hex_color(theme.primary_text_color),
    )


def add_section_heading(
    document: DocumentType, text: str, theme: ResumeTheme = DEFAULT_THEME
):
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text.upper())
    set_font(
        run,
        theme,
        size=theme.section_heading_size,
        bold=True,
        color=hex_color(theme.blue),
    )
    set_paragraph_spacing(
        paragraph,
        before=theme.section_before,
        after=theme.section_after,
        line_spacing=theme.single_line_spacing,
    )
    set_keep_with_next(paragraph)
    add_bottom_border(paragraph, color=theme.light_grey, size=theme.section_border_size)
    return paragraph


def add_body_paragraph(
    document: DocumentType,
    text: str,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    style_name: str = "Normal",
    before: float = 0,
    after: float | None = None,
    color: RGBColor | None = None,
):
    paragraph = document.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    set_font(
        run,
        theme,
        size=theme.body_size,
        color=color or hex_color(theme.primary_text_color),
    )
    set_paragraph_spacing(
        paragraph,
        before=before,
        after=theme.body_after if after is None else after,
        line_spacing=theme.body_line_spacing,
    )
    return paragraph


def add_contact_line(
    document: DocumentType,
    line: str,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    is_last: bool = False,
):
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(line)
    set_font(run, theme, size=theme.contact_size, color=hex_color(theme.grey))
    set_paragraph_spacing(
        paragraph,
        before=0,
        after=theme.contact_after if is_last else 0,
        line_spacing=theme.single_line_spacing,
    )


def add_contact_block(
    document: DocumentType, lines: Sequence[str], theme: ResumeTheme = DEFAULT_THEME
) -> None:
    for idx, line in enumerate(lines):
        add_contact_line(document, line, theme, is_last=idx == len(lines) - 1)


def add_title_block(
    document: DocumentType, meta: ResumeMeta, theme: ResumeTheme = DEFAULT_THEME
) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(meta.name)
    set_font(r, theme, size=theme.title_size, bold=True, color=hex_color(theme.ink))
    set_paragraph_spacing(
        p, before=0, after=theme.title_after, line_spacing=theme.single_line_spacing
    )
    set_keep_with_next(p)

    p = document.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(meta.title)
    set_font(r, theme, size=theme.subtitle_size, bold=True, color=hex_color(theme.blue))
    set_paragraph_spacing(
        p, before=0, after=theme.subtitle_after, line_spacing=theme.single_line_spacing
    )
    set_keep_with_next(p)

    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(meta.tagline)
    set_font(r, theme, size=theme.tagline_size, color=hex_color(theme.grey))
    set_paragraph_spacing(
        p, before=0, after=theme.tagline_after, line_spacing=theme.single_line_spacing
    )

    add_contact_block(document, meta.contact_lines, theme)


def add_skill_line(
    document: DocumentType, label: str, text: str, theme: ResumeTheme = DEFAULT_THEME
) -> None:
    p = document.add_paragraph(style="Heading 3")
    r1 = p.add_run(label)
    set_font(
        r1, theme, size=theme.skill_label_size, bold=True, color=hex_color(theme.blue)
    )
    r2 = p.add_run(f": {text}")
    set_font(
        r2,
        theme,
        size=theme.skill_label_size,
        color=hex_color(theme.primary_text_color),
    )
    set_paragraph_spacing(
        p, before=0, after=theme.skill_after, line_spacing=theme.compact_line_spacing
    )


def add_role_entry(
    document: DocumentType,
    heading_left: str,
    date_right: str,
    descriptions: Sequence[str],
    bullets: Sequence[str],
    tech: str,
    theme: ResumeTheme = DEFAULT_THEME,
    *,
    left_indent: float = 0,
    align_date: bool = True,
) -> None:
    add_entry_heading(
        document,
        heading_left,
        date_right,
        theme,
        left_indent=left_indent,
        align_date=align_date,
    )

    for idx, description in enumerate(descriptions):
        desc = document.add_paragraph(style="Normal")
        r = desc.add_run(description)
        set_font(
            r, theme, size=theme.body_size, color=hex_color(theme.primary_text_color)
        )
        is_last_description = idx == len(descriptions) - 1
        tail_after = (
            theme.role_tech_after
            if is_last_description and not bullets and not tech.strip()
            else theme.role_description_after
        )
        desc.paragraph_format.left_indent = Inches(left_indent)
        set_paragraph_spacing(
            desc,
            before=theme.role_description_before if idx == 0 else 0,
            after=tail_after,
            line_spacing=theme.role_description_line_spacing,
        )

    for bullet in bullets:
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(bullet)
        set_font(
            r, theme, size=theme.body_size, color=hex_color(theme.primary_text_color)
        )
        p.paragraph_format.left_indent = Inches(
            theme.list_bullet_left_indent + left_indent
        )
        set_paragraph_spacing(
            p,
            before=0,
            after=theme.role_bullet_after,
            line_spacing=theme.compact_line_spacing,
        )

    if tech.strip():
        tech_p = document.add_paragraph(style="Normal")
        t1 = tech_p.add_run("Tech:")
        set_font(
            t1,
            theme,
            size=theme.tech_label_size,
            italic=True,
            color=hex_color(theme.grey),
        )
        t2 = tech_p.add_run(f" {tech}")
        set_font(t2, theme, size=theme.tech_value_size, color=hex_color(theme.grey))
        tech_p.paragraph_format.left_indent = Inches(left_indent)
        set_paragraph_spacing(
            tech_p,
            before=theme.role_tech_before,
            after=theme.role_tech_after,
            line_spacing=theme.single_line_spacing,
        )


def add_nested_role_entry(
    document: DocumentType,
    heading_left: str,
    date_right: str,
    descriptions: Sequence[str],
    bullets: Sequence[str],
    tech: str,
    theme: ResumeTheme = DEFAULT_THEME,
) -> None:
    add_role_entry(
        document,
        heading_left,
        date_right,
        descriptions,
        bullets,
        tech,
        theme,
        left_indent=theme.nested_role_left_indent,
    )


def add_education_entry(
    document: DocumentType,
    heading_left: str,
    date_right: str,
    school: str,
    theme: ResumeTheme = DEFAULT_THEME,
) -> None:
    add_entry_heading(
        document,
        heading_left,
        date_right,
        theme,
        style_name="Normal",
    )

    p_school = document.add_paragraph(style="Normal")
    r_school = p_school.add_run(school)
    set_font(
        r_school, theme, size=theme.body_size, color=hex_color(theme.primary_text_color)
    )
    set_paragraph_spacing(
        p_school,
        before=2,
        after=theme.education_school_after,
        line_spacing=theme.compact_line_spacing,
    )
