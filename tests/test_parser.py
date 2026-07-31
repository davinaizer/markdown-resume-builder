from __future__ import annotations

import io
import shutil
import tempfile
import unittest
from contextlib import redirect_stderr
from pathlib import Path

import frontmatter
from docx import Document as load_docx_document
from docx.oxml.ns import qn
from frontmatter.default_handlers import YAMLHandler

from resume_builder.cli import resolve_output_path, resolve_source_path
from resume_builder.docx_utils import add_entry_heading
from resume_builder.parser import parse_resume_source
from resume_builder.renderer import build_doc_from_source, build_markdown_from_source
from resume_builder.sections import SECTION_DEFINITIONS, load_resume_directory

PROJECT_ROOT = Path(__file__).resolve().parents[1]
RESUME_SOURCE = PROJECT_ROOT / "source" / "canon-resume"


def read_frontmatter_title(path: Path) -> str:
    text = path.read_text(encoding="utf-8")
    handler = YAMLHandler()
    if not handler.detect(text):
        raise AssertionError(f"Missing YAML front matter in {path}")
    post = frontmatter.loads(text, handler=handler)
    title = post.metadata.get("title")
    if not isinstance(title, str):
        raise TypeError(f"Missing title front matter in {path}")
    return title


def read_section_titles(source: Path) -> dict[str, str]:
    titles: dict[str, str] = {}
    for definition in SECTION_DEFINITIONS:
        section_path = source / "sections" / definition.filename
        if section_path.is_file():
            titles[definition.kind] = read_frontmatter_title(section_path)
    return titles


def copy_resume_source(temp_dir: str) -> Path:
    source = Path(temp_dir) / "resume"
    shutil.copytree(RESUME_SOURCE, source)
    return source


def heading_texts(document) -> list[str]:
    return [
        paragraph.text
        for paragraph in document.paragraphs
        if getattr(paragraph.style, "name", None) == "Heading 1"
    ]


class ResumeDirectoryParserTests(unittest.TestCase):
    def test_entry_heading_uses_a_right_tab_and_non_breaking_date(self) -> None:
        document = load_docx_document()

        add_entry_heading(document, "A long role title", "11/2024 – 11/2025")
        paragraph = document.paragraphs[0]

        self.assertEqual(len(document.tables), 0)
        self.assertEqual(paragraph.text, "A long role title\t11/2024\N{NO-BREAK SPACE}–\N{NO-BREAK SPACE}11/2025")
        tabs = paragraph._p.pPr.find(qn("w:tabs"))
        self.assertEqual(tabs[0].get(qn("w:val")), "right")

    def test_canonical_source_builds_a_readable_docx_smoke_test(self) -> None:
        titles = read_section_titles(RESUME_SOURCE)

        with tempfile.TemporaryDirectory() as temp_dir:
            output = Path(temp_dir) / "resume.docx"
            stderr = io.StringIO()
            with redirect_stderr(stderr):
                document = build_doc_from_source(RESUME_SOURCE)
                document.save(str(output))

            reloaded = load_docx_document(str(output))

        self.assertEqual(stderr.getvalue(), "")
        self.assertEqual(document.core_properties.title, "Davi Naizer Santos Resume")
        self.assertEqual(document.core_properties.subject, "Resume")
        self.assertEqual(document.core_properties.author, "Davi Naizer Santos")
        self.assertEqual(
            heading_texts(reloaded),
            [
                titles["summary"].upper(),
                titles["core_skills"].upper(),
                titles["professional_experience"].upper(),
                titles["education"].upper(),
            ],
        )

    def test_cli_resolves_named_sources_and_relative_outputs_under_their_roots(
        self,
    ) -> None:
        self.assertEqual(
            resolve_source_path("canon-resume"), Path("source/canon-resume")
        )
        self.assertEqual(
            resolve_output_path("resume-test.docx"), Path("output/resume-test.docx")
        )

    def test_cli_preserves_existing_source_and_absolute_output_paths(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "resume"
            source.mkdir()
            output = Path(temp_dir) / "resume.docx"

            self.assertEqual(resolve_source_path(source), source)
            self.assertEqual(resolve_output_path(output), output)

    def test_single_file_source_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir) / "resume.md"
            source.write_text("---\nname: Test\n---\n", encoding="utf-8")

            with self.assertRaisesRegex(
                NotADirectoryError, "Resume source must be an existing directory"
            ):
                parse_resume_source(source)

    def test_missing_source_directory_is_rejected(self) -> None:
        source = Path("source/does-not-exist")

        with self.assertRaisesRegex(
            NotADirectoryError, "Resume source must be an existing directory"
        ):
            parse_resume_source(source)

    def test_parses_split_resume_source(self) -> None:
        content = parse_resume_source(RESUME_SOURCE)

        self.assertEqual(content.meta.name, "Davi Naizer Santos")
        self.assertEqual(
            {section.value for section in content.present_sections},
            {"summary", "core_skills", "professional_experience", "education"},
        )
        self.assertIsNone(content.selected_project)

    def test_loads_section_titles(self) -> None:
        source = load_resume_directory(RESUME_SOURCE)
        titles = read_section_titles(RESUME_SOURCE)

        self.assertEqual(
            [(section.kind, section.title) for section in source.sections],
            [
                ("summary", titles["summary"]),
                ("core_skills", titles["core_skills"]),
                ("professional_experience", titles["professional_experience"]),
                ("education", titles["education"]),
            ],
        )

    def test_meta_section_order_controls_rendering_order(self) -> None:
        titles = read_section_titles(RESUME_SOURCE)
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            meta_path = source / "meta.md"
            meta_markdown = meta_path.read_text(encoding="utf-8").replace(
                "sections:\n  - summary\n  - core_skills\n  - professional_experience\n  - education\n",
                "sections:\n  - education\n  - summary\n  - core_skills\n  - professional_experience\n",
                1,
            )
            meta_path.write_text(meta_markdown, encoding="utf-8")

            loaded_source = load_resume_directory(source)
            content = parse_resume_source(source)
            document = build_doc_from_source(source)
            markdown = build_markdown_from_source(source)

        self.assertEqual(
            [section.kind for section in loaded_source.sections],
            ["education", "summary", "core_skills", "professional_experience"],
        )
        self.assertEqual(
            heading_texts(document),
            [
                titles["education"].upper(),
                titles["summary"].upper(),
                titles["core_skills"].upper(),
                titles["professional_experience"].upper(),
            ],
        )
        self.assertLess(markdown.index(f"# {titles['education']}"), markdown.index(f"# {titles['summary']}"))
        self.assertLess(markdown.index(f"# {titles['summary']}"), markdown.index(f"# {titles['core_skills']}"))
        self.assertLess(markdown.index(f"# {titles['core_skills']}"), markdown.index(f"# {titles['professional_experience']}"))
        self.assertEqual(content.section_titles.education, titles["education"])

    def test_section_files_require_frontmatter(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir)
            (source / "sections").mkdir()
            (source / "meta.md").write_text("---\nname: Test\n---\n", encoding="utf-8")
            (source / "sections" / "summary.md").write_text(
                "No frontmatter", encoding="utf-8"
            )

            with self.assertRaisesRegex(
                ValueError, "must start with YAML front matter"
            ):
                load_resume_directory(source)

    def test_section_files_require_a_title(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            (source / "sections" / "summary.md").write_text(
                "---\nother: value\n---\n\nSummary.\n", encoding="utf-8"
            )

            with self.assertRaisesRegex(ValueError, "field 'title' is required"):
                load_resume_directory(source)

    def test_missing_metadata_file_fails_clearly(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = Path(temp_dir)
            (source / "sections").mkdir()

            with self.assertRaisesRegex(
                FileNotFoundError, "Resume metadata file not found"
            ):
                load_resume_directory(source)

    def test_each_missing_required_section_warns_once_and_is_not_rendered(self) -> None:
        required_definitions = [
            definition for definition in SECTION_DEFINITIONS if not definition.optional
        ]
        titles = read_section_titles(RESUME_SOURCE)

        for missing_definition in required_definitions:
            with (
                self.subTest(filename=missing_definition.filename),
                tempfile.TemporaryDirectory() as temp_dir,
            ):
                source = copy_resume_source(temp_dir)
                missing_path = source / "sections" / missing_definition.filename
                missing_path.unlink()

                stderr = io.StringIO()
                with redirect_stderr(stderr):
                    document = build_doc_from_source(source)

                warning = stderr.getvalue()
                self.assertEqual(warning.count("Warning:"), 1)
                self.assertIn(str(missing_path), warning)
                headings = heading_texts(document)
                expected_headings = [
                    titles["summary"].upper(),
                    titles["core_skills"].upper(),
                    titles["professional_experience"].upper(),
                    titles["education"].upper(),
                ]
                missing_index = {
                    "summary.md": 0,
                    "core-skills.md": 1,
                    "professional-experience.md": 2,
                    "education.md": 3,
                }[missing_definition.filename]
                expected_headings.pop(missing_index)
                self.assertEqual(headings, expected_headings)

    def test_all_missing_required_sections_warn_and_render_no_section_headings(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            required_definitions = [
                definition
                for definition in SECTION_DEFINITIONS
                if not definition.optional
            ]
            for definition in required_definitions:
                (source / "sections" / definition.filename).unlink()

            stderr = io.StringIO()
            with redirect_stderr(stderr):
                document = build_doc_from_source(source)

        warning = stderr.getvalue()
        self.assertEqual(warning.count("Warning:"), len(required_definitions))
        for definition in required_definitions:
            self.assertIn(str(source / "sections" / definition.filename), warning)
        headings = heading_texts(document)
        self.assertEqual(headings, [])

    def test_remaining_editable_title_is_preserved_after_an_earlier_omission(
        self,
    ) -> None:
        titles = read_section_titles(RESUME_SOURCE)
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            missing_path = source / "sections" / "core-skills.md"
            missing_path.unlink()
            education_path = source / "sections" / "education.md"
            education_title = read_frontmatter_title(education_path)
            education_markdown = education_path.read_text(encoding="utf-8").replace(
                f"title: {education_title}", "title: Learning & Credentials", 1
            )
            education_path.write_text(education_markdown, encoding="utf-8")

            parse_stderr = io.StringIO()
            with redirect_stderr(parse_stderr):
                content = parse_resume_source(source)
            render_stderr = io.StringIO()
            with redirect_stderr(render_stderr):
                document = build_doc_from_source(source)

        self.assertEqual(parse_stderr.getvalue().count("Warning:"), 1)
        self.assertEqual(render_stderr.getvalue().count("Warning:"), 1)
        self.assertIn(str(missing_path), parse_stderr.getvalue())
        self.assertIn(str(missing_path), render_stderr.getvalue())
        self.assertNotIn("core_skills", content.present_sections)
        self.assertEqual(content.skills, ())
        self.assertIn("summary", content.present_sections)
        self.assertIn("professional_experience", content.present_sections)
        self.assertIn("education", content.present_sections)
        self.assertEqual(content.section_titles.education, "Learning & Credentials")
        headings = heading_texts(document)
        self.assertEqual(
            headings,
            [
                titles["summary"].upper(),
                titles["professional_experience"].upper(),
                "LEARNING & CREDENTIALS",
            ],
        )

    def test_missing_optional_section_is_silent_and_not_rendered(self) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)

            stderr = io.StringIO()
            with redirect_stderr(stderr):
                loaded_source = load_resume_directory(source)
                content = parse_resume_source(source)
                document = build_doc_from_source(source)

        self.assertEqual(stderr.getvalue(), "")
        self.assertNotIn(
            "selected_project", [section.kind for section in loaded_source.sections]
        )
        self.assertNotIn("selected_project", content.present_sections)
        self.assertIsNone(content.selected_project)
        headings = heading_texts(document)
        self.assertNotIn("SELECTED PROJECT", headings)

    def test_selected_project_is_loaded_and_rendered_with_its_editable_title(
        self,
    ) -> None:
        titles = read_section_titles(RESUME_SOURCE)
        markdown = """---
title: Selected Work
---

## Project | 2026

Description.

- Result.

Tech: Python
"""
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            meta_path = source / "meta.md"
            meta_markdown = meta_path.read_text(encoding="utf-8").replace(
                "sections:\n  - summary\n  - core_skills\n  - professional_experience\n  - education\n",
                "sections:\n  - summary\n  - core_skills\n  - professional_experience\n  - selected_project\n  - education\n",
                1,
            )
            meta_path.write_text(meta_markdown, encoding="utf-8")
            (source / "sections" / "selected-project.md").write_text(
                markdown, encoding="utf-8"
            )

            loaded_source = load_resume_directory(source)
            content = parse_resume_source(source)
            document = build_doc_from_source(source)

        selected_section = next(
            section
            for section in loaded_source.sections
            if section.kind == "selected_project"
        )
        self.assertEqual(selected_section.title, "Selected Work")
        self.assertIn("selected_project", content.present_sections)
        self.assertEqual(content.section_titles.selected_project, "Selected Work")
        self.assertIsNotNone(content.selected_project)
        assert content.selected_project is not None
        self.assertEqual(content.selected_project.heading_left, "Project")
        self.assertEqual(content.selected_project.bullets, ("Result.",))
        self.assertEqual(content.selected_project.tech, "Python")
        headings = heading_texts(document)
        self.assertEqual(
            headings,
            [
                titles["summary"].upper(),
                titles["core_skills"].upper(),
                titles["professional_experience"].upper(),
                "SELECTED WORK",
                titles["education"].upper(),
            ],
        )

    def test_build_markdown_combines_sections_in_order_with_editable_titles(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            meta_path = source / "meta.md"
            meta_markdown = meta_path.read_text(encoding="utf-8").replace(
                "sections:\n  - summary\n  - core_skills\n  - professional_experience\n  - education\n",
                "sections:\n  - summary\n  - core_skills\n  - professional_experience\n  - selected_project\n  - education\n",
                1,
            )
            meta_path.write_text(meta_markdown, encoding="utf-8")
            summary_path_copy = source / "sections" / "summary.md"
            summary_markdown = summary_path_copy.read_text(encoding="utf-8").replace(
                "title: About", "title: Professional Profile", 1
            )
            summary_path_copy.write_text(summary_markdown, encoding="utf-8")
            selected_project_path = source / "sections" / "selected-project.md"
            selected_project_path.write_text(
                "---\ntitle: Selected Work\n---\n\n## Project | 2026\n\nDescription.\n\n- Result.\n\nTech: Python\n",
                encoding="utf-8",
            )

            markdown = build_markdown_from_source(source)

        self.assertIn("# Davi Naizer Santos", markdown)
        self.assertIn("**Senior Frontend Engineer**", markdown)
        self.assertIn("# Professional Profile", markdown)
        self.assertIn("# Selected Work", markdown)
        self.assertLess(markdown.index("# Professional Profile"), markdown.index("# Core Skills"))
        self.assertLess(markdown.index("# Core Skills"), markdown.index("# Professional Experience"))
        self.assertLess(markdown.index("# Professional Experience"), markdown.index("# Selected Work"))
        self.assertLess(markdown.index("# Selected Work"), markdown.index("# Education"))

    def test_editable_title_is_passed_to_model_and_renderer_without_changing_section_type(
        self,
    ) -> None:
        titles = read_section_titles(RESUME_SOURCE)
        with tempfile.TemporaryDirectory() as temp_dir:
            source = copy_resume_source(temp_dir)
            summary_path_copy = source / "sections" / "summary.md"
            summary_markdown = summary_path_copy.read_text(encoding="utf-8").replace(
                f"title: {titles['summary']}", "title: Professional Profile", 1
            )
            summary_path_copy.write_text(summary_markdown, encoding="utf-8")

            content = parse_resume_source(source)
            document = build_doc_from_source(source)

        self.assertEqual(content.section_titles.summary, "Professional Profile")
        headings = heading_texts(document)
        self.assertEqual(
            headings,
            [
                "PROFESSIONAL PROFILE",
                titles["core_skills"].upper(),
                titles["professional_experience"].upper(),
                titles["education"].upper(),
            ],
        )


if __name__ == "__main__":
    unittest.main()
