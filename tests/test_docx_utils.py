from __future__ import annotations

import io
import unittest
from contextlib import redirect_stderr

from docx import Document as load_docx_document
from docx.oxml.ns import qn

from resume_builder.docx_utils import (
    add_entry_heading,
    add_nested_role_entry,
    add_role_entry,
)
from resume_builder.parser import parse_experience_lines


class DocxUtilsTests(unittest.TestCase):
    def test_entry_heading_uses_a_right_tab_and_non_breaking_date(self) -> None:
        document = load_docx_document()

        add_entry_heading(document, "A long role title", "11/2024 – 11/2025")
        paragraph = document.paragraphs[0]

        self.assertEqual(len(document.tables), 0)
        self.assertEqual(
            paragraph.text,
            "A long role title\t11/2024\N{NO-BREAK SPACE}–\N{NO-BREAK SPACE}11/2025",
        )
        paragraph_properties = paragraph._p.pPr
        self.assertIsNotNone(paragraph_properties)
        assert paragraph_properties is not None
        tabs = paragraph_properties.find(qn("w:tabs"))
        self.assertIsNotNone(tabs)
        assert tabs is not None
        self.assertEqual(tabs[0].get(qn("w:val")), "right")

    def test_role_entries_preserve_and_render_every_introductory_paragraph(
        self,
    ) -> None:
        lines = [
            "## Role | Company | 2024 – Present",
            "<!-- experience: employment -->",
            "",
            "First paragraph.",
            "",
            "Second paragraph.",
            "",
            "Third paragraph.",
            "",
            "- Result.",
            "",
            "**Tech:** Python",
        ]

        stderr = io.StringIO()
        with redirect_stderr(stderr):
            entries = parse_experience_lines(lines)

        self.assertEqual(stderr.getvalue(), "")
        self.assertEqual(
            entries[0].descriptions,
            ("First paragraph.", "Second paragraph.", "Third paragraph."),
        )

        document = load_docx_document()
        add_role_entry(
            document,
            "Role | Company",
            entries[0].date_right,
            entries[0].descriptions,
            entries[0].bullets,
            entries[0].tech,
        )
        rendered_text = [paragraph.text for paragraph in document.paragraphs]
        for description in entries[0].descriptions:
            self.assertIn(description, rendered_text)

    def test_nested_role_entry_is_indented_and_ordered_after_parent(self) -> None:
        document = load_docx_document()
        add_role_entry(
            document, "Company | London", "2020 – 2024", ("Parent.",), (), ""
        )
        add_nested_role_entry(
            document,
            "Engineer | Brand",
            "2023 – 2024",
            ("Role one.", "Role two."),
            ("Delivered.",),
            "Python",
        )

        paragraphs = document.paragraphs
        texts = [paragraph.text for paragraph in paragraphs]
        parent_heading_index = texts.index(
            "Company | London\t2020\N{NO-BREAK SPACE}–\N{NO-BREAK SPACE}2024"
        )
        nested_heading_index = texts.index(
            "Engineer | Brand\t2023\N{NO-BREAK SPACE}–\N{NO-BREAK SPACE}2024"
        )
        nested_content = ["Role one.", "Role two.", "Delivered.", "Tech: Python"]
        self.assertLess(parent_heading_index, nested_heading_index)
        self.assertEqual(
            texts[nested_heading_index + 1 : nested_heading_index + 5], nested_content
        )
        nested_heading = paragraphs[nested_heading_index]
        self.assertGreater(nested_heading.paragraph_format.left_indent.pt, 0)
