from __future__ import annotations

import argparse
from dataclasses import dataclass
from pathlib import Path

import frontmatter
from frontmatter.default_handlers import YAMLHandler
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from tools.resume_theme import DEFAULT_THEME, ResumeTheme


OUT_PATH = Path("docs/my-resume.docx")


@dataclass
class ResumeMeta:
    name: str
    title: str
    tagline: str
    contact_lines: list[str]


@dataclass
class SkillLine:
    label: str
    value: str


@dataclass
class EntryBlock:
    heading_left: str
    date_right: str
    description: str
    bullets: list[str]
    tech: str


@dataclass
class EducationBlock:
    heading_left: str
    date_right: str
    school: str


@dataclass
class ResumeContent:
    meta: ResumeMeta
    summary: list[str]
    skills: list[SkillLine]
    experience: list[EntryBlock]
    selected_project: EntryBlock
    education: list[EducationBlock]


def clean_md_text(text: str) -> str:
    text = text.strip()
    if text.startswith("## "):
        text = text[3:]
    elif text.startswith("# "):
        text = text[2:]
    text = text.replace("**", "")
    return text.strip()


def is_rule(line: str) -> bool:
    return line.strip() == "---"


def is_h1(line: str) -> bool:
    return line.lstrip().startswith("# ") and not line.lstrip().startswith("## ")


def is_h2(line: str) -> bool:
    return line.lstrip().startswith("## ")


def is_bullet(line: str) -> bool:
    return line.lstrip().startswith("- ")


def split_heading_date(text: str) -> tuple[str, str]:
    parts = [part.strip() for part in text.split("|")]
    if len(parts) < 2:
        raise ValueError(f"Expected heading with date separated by '|': {text}")
    left = " | ".join(parts[:-1]).strip()
    right = parts[-1].strip()
    return left, right


def next_content_line(lines: list[str], start: int) -> tuple[str | None, int]:
    i = start
    while i < len(lines):
        line = lines[i].strip()
        if line and not is_rule(line):
            return line, i
        i += 1
    return None, len(lines)


def hex_color(value: str) -> RGBColor:
    value = value.strip().lstrip("#")
    if len(value) != 6:
        raise ValueError(f"Expected 6-digit hex color, got {value!r}")
    return RGBColor(int(value[0:2], 16), int(value[2:4], 16), int(value[4:6], 16))


def _require_string(metadata: dict, key: str) -> str:
    value = metadata.get(key)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"Front matter field '{key}' is required and must be a non-empty string")
    return clean_md_text(value)


def _require_string_list(metadata: dict, key: str) -> list[str]:
    value = metadata.get(key)
    if not isinstance(value, list) or not value:
        raise ValueError(f"Front matter field '{key}' is required and must be a non-empty list of strings")
    cleaned: list[str] = []
    for item in value:
        if not isinstance(item, str) or not item.strip():
            raise ValueError(f"Front matter field '{key}' must contain only non-empty strings")
        cleaned.append(item.strip())
    return cleaned


def parse_resume_markdown(path: Path) -> ResumeContent:
    text = path.read_text(encoding="utf-8")
    handler = YAMLHandler()
    if not handler.detect(text):
        raise ValueError("Resume markdown must start with YAML front matter")

    post = frontmatter.loads(text, handler=handler)
    metadata = post.metadata
    name = _require_string(metadata, "name")
    title = _require_string(metadata, "title")
    tagline = _require_string(metadata, "tagline")
    contact_lines = _require_string_list(metadata, "contact_lines")

    lines = post.content.splitlines()
    i = 0

    summary: list[str] = []
    skills: list[SkillLine] = []
    experience: list[EntryBlock] = []
    selected_project: EntryBlock | None = None
    education: list[EducationBlock] = []

    def collect_paragraph() -> str:
        nonlocal i
        line, idx = next_content_line(lines, i)
        if line is None:
            raise ValueError("Expected paragraph content")
        i = idx + 1
        return clean_md_text(line)

    while i < len(lines):
        line = lines[i].strip()
        if not line or is_rule(line):
            i += 1
            continue
        if not is_h1(line):
            i += 1
            continue

        section = clean_md_text(line)
        i += 1

        if section == "Summary":
            while i < len(lines) and not is_h1(lines[i]):
                line = lines[i].strip()
                if line and not is_rule(line):
                    summary.append(clean_md_text(line))
                i += 1
        elif section == "Core Skills":
            while i < len(lines) and not is_h1(lines[i]):
                line = lines[i].strip()
                if is_h2(line):
                    label = clean_md_text(line)
                    if label.startswith("## "):
                        label = label[3:]
                    i += 1
                    value = collect_paragraph()
                    skills.append(SkillLine(label=label, value=value))
                else:
                    i += 1
        elif section == "Professional Experience":
            while i < len(lines) and not is_h1(lines[i]):
                line = lines[i].strip()
                if is_h2(line):
                    heading = clean_md_text(line)
                    if heading.startswith("## "):
                        heading = heading[3:]
                    heading_left, date_right = split_heading_date(heading)
                    i += 1
                    description = collect_paragraph()
                    bullets: list[str] = []
                    tech = ""
                    while i < len(lines):
                        peek = lines[i].strip()
                        if not peek:
                            i += 1
                            continue
                        if is_rule(peek) or is_h2(peek) or is_h1(peek):
                            break
                        if is_bullet(peek):
                            bullets.append(clean_md_text(peek[2:]))
                            i += 1
                            continue
                        if peek.startswith("Tech:"):
                            tech = clean_md_text(peek)[5:].strip()
                            i += 1
                            break
                        i += 1
                    experience.append(
                        EntryBlock(
                            heading_left=heading_left,
                            date_right=date_right,
                            description=description,
                            bullets=bullets,
                            tech=tech,
                        )
                    )
                else:
                    i += 1
        elif section == "Selected Project":
            while i < len(lines) and not is_h1(lines[i]):
                line = lines[i].strip()
                if is_h2(line):
                    heading = clean_md_text(line)
                    if heading.startswith("## "):
                        heading = heading[3:]
                    heading_left, date_right = split_heading_date(heading)
                    i += 1
                    description = collect_paragraph()
                    bullets = []
                    tech = ""
                    while i < len(lines):
                        peek = lines[i].strip()
                        if not peek:
                            i += 1
                            continue
                        if is_rule(peek) or is_h2(peek) or is_h1(peek):
                            break
                        if is_bullet(peek):
                            bullets.append(clean_md_text(peek[2:]))
                            i += 1
                            continue
                        if peek.startswith("Tech:"):
                            tech = clean_md_text(peek)[5:].strip()
                            i += 1
                            break
                        i += 1
                    selected_project = EntryBlock(
                        heading_left=heading_left,
                        date_right=date_right,
                        description=description,
                        bullets=bullets,
                        tech=tech,
                    )
                else:
                    i += 1
        elif section == "Education":
            while i < len(lines) and not is_h1(lines[i]):
                line = lines[i].strip()
                if is_h2(line):
                    heading = clean_md_text(line)
                    if heading.startswith("## "):
                        heading = heading[3:]
                    heading_left, date_right = split_heading_date(heading)
                    i += 1
                    school = collect_paragraph()
                    education.append(
                        EducationBlock(
                            heading_left=heading_left,
                            date_right=date_right,
                            school=school,
                        )
                    )
                else:
                    i += 1
        else:
            while i < len(lines) and not is_h1(lines[i]):
                i += 1

    if selected_project is None:
        raise ValueError("Selected Project section was not parsed")

    return ResumeContent(
        meta=ResumeMeta(name=name, title=title, tagline=tagline, contact_lines=contact_lines),
        summary=summary,
        skills=skills,
        experience=experience,
        selected_project=selected_project,
        education=education,
    )


def set_font(run, theme: ResumeTheme = DEFAULT_THEME, *, name: str | None = None,
             size: float | None = None, bold: bool = False, italic: bool = False,
             color: RGBColor | None = None) -> None:
    name = name or theme.font_family
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.bold = bold
    run.font.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color


def set_style_font(style, theme: ResumeTheme = DEFAULT_THEME, *, name: str | None = None,
                   size: float = 11, bold: bool = False,
                   color: RGBColor | None = None, italic: bool = False) -> None:
    name = name or theme.font_family
    style.font.name = name
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    if color is not None:
        style.font.color.rgb = color


def configure_paragraph_style(style, theme: ResumeTheme = DEFAULT_THEME, *, size: float,
                              bold: bool = False, color: RGBColor | None = None,
                              italic: bool = False, before: float = 0,
                              after: float = 0, line_spacing: float = 1.2,
                              left_indent: Inches | None = None,
                              first_line_indent: Inches | None = None,
                              alignment: WD_ALIGN_PARAGRAPH | None = None) -> None:
    set_style_font(style, theme, size=size, bold=bold, color=color, italic=italic)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    style.paragraph_format.line_spacing = line_spacing
    if left_indent is not None:
        style.paragraph_format.left_indent = left_indent
    if first_line_indent is not None:
        style.paragraph_format.first_line_indent = first_line_indent
    if alignment is not None:
        style.paragraph_format.alignment = alignment


def set_paragraph_spacing(paragraph, *, before: float = 0, after: float = 0,
                          line_spacing: float = 1.2) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    fmt.line_spacing = line_spacing


def set_keep_with_next(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepNext")
    ppr.append(keep)


def set_keep_lines(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    keep = OxmlElement("w:keepLines")
    ppr.append(keep)


def add_bottom_border(paragraph, color: str = DEFAULT_THEME.light_grey, size: str = DEFAULT_THEME.section_border_size) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    p_bdr = ppr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        ppr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def clear_document(document: Document) -> None:
    body = document._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_table_fixed_width(table, widths: list[Inches]) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for idx, width in enumerate(widths):
        if idx < len(grid.gridCol_lst):
            grid.gridCol_lst[idx].set(qn("w:w"), str(int(width.inches * 1440)))
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(int(widths[idx].inches * 1440)))


def remove_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_cell_margins(cell, *, top=80, start=80, bottom=80, end=80) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def add_section_heading(document: Document, text: str, theme: ResumeTheme = DEFAULT_THEME):
    paragraph = document.add_paragraph(style="Heading 1")
    run = paragraph.add_run(text.upper())
    set_font(run, theme, size=theme.section_heading_size, bold=True, color=hex_color(theme.blue))
    set_paragraph_spacing(paragraph, before=theme.section_before, after=theme.section_after, line_spacing=theme.single_line_spacing)
    set_keep_with_next(paragraph)
    add_bottom_border(paragraph, color=theme.light_grey, size=theme.section_border_size)
    return paragraph


def add_body_paragraph(document: Document, text: str, theme: ResumeTheme = DEFAULT_THEME,
                       *, style_name: str = "Normal", before: float = 0, after: float | None = None,
                       color: RGBColor | None = None):
    paragraph = document.add_paragraph(style=style_name)
    run = paragraph.add_run(text)
    set_font(run, theme, size=theme.body_size, color=color or hex_color(theme.primary_text_color))
    set_paragraph_spacing(paragraph, before=before, after=theme.body_after if after is None else after, line_spacing=theme.body_line_spacing)
    return paragraph


def add_contact_line(document: Document, line: str, theme: ResumeTheme = DEFAULT_THEME, *, is_last: bool = False):
    paragraph = document.add_paragraph(style="Normal")
    run = paragraph.add_run(line)
    set_font(run, theme, size=theme.contact_size, color=hex_color(theme.grey))
    set_paragraph_spacing(
        paragraph,
        before=0,
        after=theme.contact_after if is_last else 0,
        line_spacing=theme.single_line_spacing,
    )


def add_contact_block(document: Document, lines: list[str], theme: ResumeTheme = DEFAULT_THEME) -> None:
    for idx, line in enumerate(lines):
        add_contact_line(document, line, theme, is_last=idx == len(lines) - 1)


def add_title_block(document: Document, meta: ResumeMeta, theme: ResumeTheme = DEFAULT_THEME) -> None:
    p = document.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(meta.name)
    set_font(r, theme, size=theme.title_size, bold=True, color=hex_color(theme.ink))
    set_paragraph_spacing(p, before=0, after=theme.title_after, line_spacing=theme.single_line_spacing)
    set_keep_with_next(p)

    p = document.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(meta.title)
    set_font(r, theme, size=theme.subtitle_size, bold=True, color=hex_color(theme.blue))
    set_paragraph_spacing(p, before=0, after=theme.subtitle_after, line_spacing=theme.single_line_spacing)
    set_keep_with_next(p)

    p = document.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(meta.tagline)
    set_font(r, theme, size=theme.tagline_size, color=hex_color(theme.grey))
    set_paragraph_spacing(p, before=0, after=theme.tagline_after, line_spacing=theme.single_line_spacing)

    add_contact_block(document, meta.contact_lines, theme)


def add_skill_line(document: Document, label: str, text: str, theme: ResumeTheme = DEFAULT_THEME) -> None:
    p = document.add_paragraph(style="Heading 3")
    r1 = p.add_run(label)
    set_font(r1, theme, size=theme.skill_label_size, bold=True, color=hex_color(theme.blue))
    r2 = p.add_run(f": {text}")
    set_font(r2, theme, size=theme.skill_label_size, color=hex_color(theme.primary_text_color))
    set_paragraph_spacing(p, before=0, after=theme.skill_after, line_spacing=theme.compact_line_spacing)


def add_role_entry(document: Document, heading_left: str, date_right: str, description: str,
                   bullets: list[str], tech: str, theme: ResumeTheme = DEFAULT_THEME) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_fixed_width(
        table,
        [Inches(theme.role_table_left_width), Inches(theme.role_table_right_width)],
    )
    remove_table_borders(table)

    left, right = table.rows[0].cells
    set_cell_margins(left, top=0, start=0, bottom=0, end=0)
    set_cell_margins(right, top=0, start=0, bottom=0, end=0)

    p_left = left.paragraphs[0]
    p_left.style = document.styles["Heading 2"]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_left.paragraph_format.space_before = Pt(0)
    p_left.paragraph_format.space_after = Pt(0)
    p_left.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r_left = p_left.add_run(heading_left)
    set_font(r_left, theme, size=theme.role_heading_size, bold=True, color=hex_color(theme.primary_text_color))

    p_right = right.paragraphs[0]
    p_right.style = document.styles["Normal"]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_right.paragraph_format.space_before = Pt(0)
    p_right.paragraph_format.space_after = Pt(0)
    p_right.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    r_right = p_right.add_run(date_right)
    set_font(r_right, theme, size=theme.date_size, color=hex_color(theme.primary_text_color))

    desc = document.add_paragraph(style="Normal")
    r = desc.add_run(description)
    set_font(r, theme, size=theme.body_size, color=hex_color(theme.primary_text_color))
    tail_after = theme.role_tech_after if not bullets and not tech.strip() else theme.role_description_after
    set_paragraph_spacing(
        desc,
        before=theme.role_description_before,
        after=tail_after,
        line_spacing=theme.role_description_line_spacing,
    )

    for bullet in bullets:
        p = document.add_paragraph(style="List Bullet")
        r = p.add_run(bullet)
        set_font(r, theme, size=theme.body_size, color=hex_color(theme.primary_text_color))
        set_paragraph_spacing(p, before=0, after=theme.role_bullet_after, line_spacing=theme.compact_line_spacing)

    if tech.strip():
        tech_p = document.add_paragraph(style="Normal")
        t1 = tech_p.add_run("Tech:")
        set_font(t1, theme, size=theme.tech_label_size, italic=True, color=hex_color(theme.grey))
        t2 = tech_p.add_run(f" {tech}")
        set_font(t2, theme, size=theme.tech_value_size, color=hex_color(theme.grey))
        set_paragraph_spacing(tech_p, before=theme.role_tech_before, after=theme.role_tech_after, line_spacing=theme.single_line_spacing)


def add_education_entry(document: Document, heading_left: str, date_right: str, school: str,
                        theme: ResumeTheme = DEFAULT_THEME) -> None:
    table = document.add_table(rows=1, cols=2)
    set_table_fixed_width(
        table,
        [Inches(theme.role_table_left_width), Inches(theme.role_table_right_width)],
    )
    remove_table_borders(table)
    left, right = table.rows[0].cells
    set_cell_margins(left, top=0, start=0, bottom=0, end=0)
    set_cell_margins(right, top=0, start=0, bottom=0, end=0)

    p_left = left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r_left = p_left.add_run(heading_left)
    set_font(r_left, theme, size=theme.role_heading_size, bold=True, color=hex_color(theme.primary_text_color))
    p_left.paragraph_format.space_after = Pt(0)
    p_left.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p_right = right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_right = p_right.add_run(date_right)
    set_font(r_right, theme, size=theme.date_size, color=hex_color(theme.primary_text_color))
    p_right.paragraph_format.space_after = Pt(0)
    p_right.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    p_school = document.add_paragraph(style="Normal")
    r_school = p_school.add_run(school)
    set_font(r_school, theme, size=theme.body_size, color=hex_color(theme.primary_text_color))
    set_paragraph_spacing(p_school, before=2, after=theme.education_school_after, line_spacing=theme.compact_line_spacing)


def add_experience_entry(document: Document, heading_left: str, date_right: str, description: str,
                         bullets: list[str], tech: str, theme: ResumeTheme = DEFAULT_THEME) -> None:
    add_role_entry(document, heading_left, date_right, description, bullets, tech, theme)


def build_doc_from_markdown(md_path: Path, theme: ResumeTheme = DEFAULT_THEME) -> Document:
    content = parse_resume_markdown(md_path)
    doc = Document()
    clear_document(doc)

    section = doc.sections[0]
    section.page_width = Inches(theme.page_width)
    section.page_height = Inches(theme.page_height)
    margin = Inches(theme.margin)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin
    section.header_distance = Inches(theme.header_distance)
    section.footer_distance = Inches(theme.footer_distance)

    core_props = doc.core_properties
    core_props.title = f"{content.meta.name} Resume"
    core_props.subject = "Resume"
    core_props.author = content.meta.name

    normal = doc.styles["Normal"]
    set_style_font(normal, theme, size=theme.body_size, color=hex_color(theme.primary_text_color))
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = theme.body_line_spacing

    configure_paragraph_style(
        doc.styles["Title"],
        theme,
        size=theme.title_size,
        bold=True,
        color=hex_color(theme.ink),
        before=0,
        after=theme.title_after,
        line_spacing=theme.single_line_spacing,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    configure_paragraph_style(
        doc.styles["Subtitle"],
        theme,
        size=theme.subtitle_size,
        bold=True,
        color=hex_color(theme.blue),
        before=0,
        after=theme.subtitle_after,
        line_spacing=theme.single_line_spacing,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    configure_paragraph_style(
        doc.styles["Heading 1"],
        theme,
        size=theme.section_heading_size,
        bold=True,
        color=hex_color(theme.blue),
        before=theme.section_before,
        after=theme.section_after,
        line_spacing=theme.single_line_spacing,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    configure_paragraph_style(
        doc.styles["Heading 2"],
        theme,
        size=theme.heading2_size,
        bold=True,
        color=hex_color(theme.primary_text_color),
        before=0,
        after=theme.heading2_after,
        line_spacing=theme.heading2_line_spacing,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    configure_paragraph_style(
        doc.styles["Heading 3"],
        theme,
        size=theme.skill_label_size,
        bold=True,
        color=hex_color(theme.blue),
        before=0,
        after=theme.skill_after,
        line_spacing=theme.compact_line_spacing,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    configure_paragraph_style(
        doc.styles["Normal"],
        theme,
        size=theme.body_size,
        color=hex_color(theme.primary_text_color),
        before=0,
        after=0,
        line_spacing=theme.body_line_spacing,
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )
    configure_paragraph_style(
        doc.styles["List Bullet"],
        theme,
        size=theme.body_size,
        color=hex_color(theme.primary_text_color),
        before=0,
        after=0,
        line_spacing=theme.compact_line_spacing,
        left_indent=Inches(theme.list_bullet_left_indent),
        first_line_indent=Inches(theme.list_bullet_first_line_indent),
        alignment=WD_ALIGN_PARAGRAPH.LEFT,
    )

    add_title_block(doc, content.meta, theme)

    add_section_heading(doc, "Summary", theme)
    for idx, paragraph in enumerate(content.summary):
        add_body_paragraph(
            doc,
            paragraph,
            theme,
            after=theme.summary_after if idx < len(content.summary) - 1 else theme.summary_last_after,
        )

    add_section_heading(doc, "Core Skills", theme)
    for skill in content.skills:
        add_skill_line(doc, skill.label, skill.value, theme)

    add_section_heading(doc, "Professional Experience", theme)
    for entry in content.experience:
        add_experience_entry(
            doc,
            entry.heading_left,
            entry.date_right,
            entry.description,
            entry.bullets,
            entry.tech,
            theme,
        )

    add_section_heading(doc, "Selected Project", theme)
    add_experience_entry(
        doc,
        content.selected_project.heading_left,
        content.selected_project.date_right,
        content.selected_project.description,
        content.selected_project.bullets,
        content.selected_project.tech,
        theme,
    )

    add_section_heading(doc, "Education", theme)
    for item in content.education:
        add_education_entry(doc, item.heading_left, item.date_right, item.school, theme)

    return doc


def main() -> None:
    parser = argparse.ArgumentParser(description="Build a resume DOCX from Markdown.")
    parser.add_argument("input_md", nargs="?", default=str(Path("docs/my-resume.md")))
    parser.add_argument("-o", "--output", default=str(OUT_PATH))
    args = parser.parse_args()

    input_md = Path(args.input_md)
    output_docx = Path(args.output)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc = build_doc_from_markdown(input_md, DEFAULT_THEME)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    print(f"Wrote {output_docx}")


if __name__ == "__main__":
    main()
